"""
Grading Tool - Parse rubrics and grade submissions.
"""
import os
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document


class GradingTool:
    """Tool for parsing rubrics and grading submissions."""
    
    def __init__(self, rubrics_directory: str = None):
        self.rubrics_directory = rubrics_directory or os.path.dirname(__file__).replace("tools", "")
        self._rubrics_cache: Dict[str, str] = {}
    
    def list_rubrics(self) -> str:
        """List available rubric files."""
        rubrics = []
        for file in os.listdir(self.rubrics_directory):
            if file.lower().endswith(('.docx', '.pdf')) and 'rubric' in file.lower():
                rubrics.append(file)
        if not rubrics:
            return "No rubric files found in the rubrics directory."
        return "Available rubrics:\n" + "\n".join(f"  - {r}" for r in rubrics)
    
    def parse_docx_rubric(self, rubric_path: str) -> str:
        """Parse a DOCX rubric file and extract grading criteria."""
        try:
            doc = Document(rubric_path)
            content = []
            
            # Extract all text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            # Extract text from tables (rubrics often use tables)
            for table in doc.tables:
                content.append("\n[TABLE]")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content.append(row_text)
                content.append("[/TABLE]\n")
            
            return "\n".join(content)
        except Exception as e:
            return f"Error parsing rubric: {e}"
    
    def load_rubric(self, rubric_name: str) -> str:
        """Load a rubric by name (searches for matching files)."""
        # Check cache first
        if rubric_name in self._rubrics_cache:
            return self._rubrics_cache[rubric_name]
        
        # Search for rubric file
        search_terms = rubric_name.lower().split()
        for file in os.listdir(self.rubrics_directory):
            if not file.lower().endswith('.docx'):
                continue
            file_lower = file.lower()
            if all(term in file_lower for term in search_terms):
                rubric_path = os.path.join(self.rubrics_directory, file)
                content = self.parse_docx_rubric(rubric_path)
                self._rubrics_cache[rubric_name] = content
                return f"Rubric: {file}\n\n{content}"
        
        return f"Could not find rubric matching '{rubric_name}'. Use list_rubrics to see available rubrics."
    
    def read_submission(self, submission_path: str) -> str:
        """Read a student submission (supports txt, docx, pdf)."""
        try:
            ext = os.path.splitext(submission_path)[1].lower()
            
            if ext == '.txt':
                with open(submission_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif ext == '.docx':
                doc = Document(submission_path)
                return "\n".join(para.text for para in doc.paragraphs)
            
            elif ext == '.pdf':
                # Basic PDF support would require PyPDF2 or similar
                return "PDF support not yet implemented. Please convert to DOCX or TXT."
            
            else:
                return f"Unsupported file format: {ext}"
        
        except Exception as e:
            return f"Error reading submission: {e}"
    
    def grade_submission(self, submission_path: str, rubric_name: str) -> str:
        """
        Grade a submission against a rubric.
        Returns the rubric and submission for the LLM to evaluate.
        """
        rubric_content = self.load_rubric(rubric_name)
        submission_content = self.read_submission(submission_path)
        
        return f"""=== RUBRIC ===
{rubric_content}

=== STUDENT SUBMISSION ===
{submission_content}

=== GRADING INSTRUCTIONS ===
Please evaluate the student submission against the rubric criteria above.
For each criterion, assign a score and provide specific feedback.
Calculate the total score and provide an overall assessment."""


# Singleton instance
_grading_tool: Optional[GradingTool] = None


def get_grading(rubrics_dir: str = None) -> GradingTool:
    """Get grading tool instance."""
    global _grading_tool
    if _grading_tool is None:
        _grading_tool = GradingTool(rubrics_dir)
    return _grading_tool
